from fastapi import FastAPI, HTTPException, Depends, Request, UploadFile, File
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
import psycopg2
import psycopg2.extras
import re
import logging
import traceback
import base64
import json
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
from openai import OpenAI
import anthropic
import mammoth
from typing import List, Optional
import os
import io
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

load_dotenv()

logging.basicConfig(
    filename=os.path.join(os.path.dirname(__file__), "logs/error.log"),
    level=logging.ERROR,
    format="%(asctime)s %(levelname)s %(message)s"
)

DEEPSEEK_API_KEY = os.environ["DEEPSEEK_API_KEY"]
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

SQL_SYSTEM_PROMPT = """你是一个求职数据分析助手，专门帮助用户查询和分析他们的求职申请记录数据库。

## 职责范围
你只能回答与以下内容相关的问题：
- 用户的求职申请记录（公司、职位、申请日期、地点、工作类型、反馈结果）
- 数据统计与分析（投递数量、地点分布、通过率、时间趋势等）
- 工作许可信息（work_permits 表：国家、签证类型、薪资门槛、永居年限）

## 拒绝规则
如果用户的问题与上述职责范围无关（例如：写代码、翻译、天气、通用知识、聊天等），
必须严格执行以下要求：
- 不得生成任何 SQL 语句
- 不得调用自身知识直接回答问题
- 不得尝试用任何方式帮助用户完成与求职数据无关的请求
- Output ONLY the following fixed reply in the same language as the user's message, nothing else:
  - If the user's message is in English: "Sorry, I can only help you analyze your job application data. Please ask questions related to your applications, e.g. 'How many companies have I applied to?' or 'Which location has the most applications?'"
  - If the user's message is in Chinese: "抱歉，我只能帮你分析求职申请数据。请提问与你的投递记录相关的问题，例如：'我投了多少家公司？' 或 '哪个地点投递最多？'"

## 数据库结构

表1：job_applications
- id：主键
- company：公司名称
- position：职位名称
- applied_date：投递日期（DATE 类型，格式 YYYY-MM-DD，年份为 2026）
- location：国家/地区（如 "Norway"、"Netherlands"）
- link：职位链接
- feedback：反馈结果（NULL=待回复，"Fail"=拒绝，"Offer"=录用，"Interview"=面试，"Online Assessment"=线上笔试）
- work_type：工作类型（Remote / Onsite / Hybrid）
- user_id：用户 ID

表2：work_permits
- country：国家
- visa：签证/工作许可类型
- annual_salary：年薪门槛（文本）
- permanent_residence：永居申请年限

## 字段映射
- "地点" / "location" / "country" → job_applications 表的 location 字段
- "没有反馈" / "pending" / "待回复" → feedback IS NULL

## SQL 生成规则
确认问题与求职数据相关后，生成标准 PostgreSQL SELECT 语句：
- 只生成 SELECT 语句，不生成 INSERT / UPDATE / DELETE
- 查询 job_applications 时必须加上 WHERE user_id = {user_id}
- 涉及 location 字段时，必须附加过滤条件：location IS NOT NULL AND location != '' AND location != 'NaN'
- 涉及"最多/最少/前N名/排名"等问题时，必须使用 GROUP BY + ORDER BY + LIMIT
- 只输出原始 SQL 语句本身，不加任何解释、不加 markdown、不加代码块"""

EXPLAIN_SYSTEM_PROMPT_EN = """You are a job-search data analyst. Answer the user's question in natural language based on the database query results.

CRITICAL: Respond in ENGLISH ONLY. Do not output any Chinese characters under any circumstances.

Requirements:
- Be concise and direct — lead with the conclusion
- If the result set is empty, tell the user there are no matching records
- Do not repeat raw data; summarize it naturally
- Keep the answer to 3-5 sentences"""

EXPLAIN_SYSTEM_PROMPT_ZH = """你是一个求职数据分析助手，根据数据库查询结果用自然语言回答用户的问题。

严格要求：只用中文回答，不得使用英文。

- 语言简洁清晰，直接给出结论
- 如果数据为空，告知用户暂无相关记录
- 不要重复展示原始数据，用自然语言总结
- 回答控制在3-5句话以内"""

def _is_english(text: str) -> bool:
    chinese = sum(1 for c in text if '一' <= c <= '鿿')
    return chinese < len(text) * 0.1

limiter = Limiter(key_func=get_remote_address)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

_origins_env = os.environ.get("ALLOWED_ORIGINS", "")
ALLOWED_ORIGINS = [o.strip() for o in _origins_env.split(",") if o.strip()]
if not ALLOWED_ORIGINS:
    import warnings
    warnings.warn("ALLOWED_ORIGINS is not set — all CORS requests will be blocked");

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE"],
    allow_headers=["Authorization", "Content-Type"],
)

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logging.error(f"{request.method} {request.url.path}\n{traceback.format_exc()}")
    return JSONResponse(status_code=500, content={"detail": "服务器内部错误，请稍后重试"})

DB_CONFIG = {
    "host": os.environ.get("DB_HOST", "localhost"),
    "port": int(os.environ.get("DB_PORT", 5432)),
    "user": os.environ.get("DB_USER", "postgres"),
    "password": os.environ["DB_PASSWORD"],
    "database": os.environ.get("DB_NAME", "jobsdb"),
}

SECRET_KEY = os.environ["SECRET_KEY"]
ALGORITHM = "HS256"
TOKEN_EXPIRE_DAYS = 30
CHAT_DAILY_LIMIT = 50

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
bearer = HTTPBearer()

BLOCKED = re.compile(
    r"\b(INSERT|UPDATE|DELETE|DROP|TRUNCATE|ALTER|CREATE|GRANT|REVOKE)\b",
    re.IGNORECASE
)
ALLOWED_TABLES = {"job_applications", "work_permits"}
_TABLE_REF = re.compile(r"\b(?:FROM|JOIN)\s+([a-zA-Z_][a-zA-Z0-9_]*)", re.IGNORECASE)

def validate_chat_sql(sql: str, user_id: int) -> str | None:
    """Return an error string if the SQL is unsafe, None if it passes."""
    if ";" in sql:
        return "不允许多语句查询"
    if BLOCKED.search(sql):
        return "包含不允许的操作"
    referenced = {m.group(1).lower() for m in _TABLE_REF.finditer(sql)}
    disallowed = referenced - ALLOWED_TABLES
    if disallowed:
        return f"不允许查询的表：{disallowed}"
    if "job_applications" in referenced:
        if not re.search(rf"\buser_id\s*=\s*{user_id}\b", sql):
            return "缺少 user_id 过滤条件"
    return None

# ── Models ──

class ApplicationRequest(BaseModel):
    company: str
    position: str
    applied_date: str | None = None
    location: str | None = None
    link: str | None = None
    feedback: str | None = None
    work_type: str | None = None
    notes: str | None = None

class AuthRequest(BaseModel):
    email: str
    password: str
    invite_code: Optional[str] = None

class ResetPasswordRequest(BaseModel):
    new_password: str

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str = Field(min_length=6)

class FeedbackRequest(BaseModel):
    category: str = Field(default="other", max_length=50)
    content: str = Field(min_length=1, max_length=2000)

class ChatMessage(BaseModel):
    role: str
    content: str

class AnalyzeRequest(BaseModel):
    message: str
    type: str = ""

class ExportRequest(BaseModel):
    content: str
    format: str  # "docx" or "pdf"

class ChatRequest(BaseModel):
    message: str = Field(max_length=500)
    history: Optional[List[ChatMessage]] = []

# ── Auth helpers ──

def get_db():
    return psycopg2.connect(**DB_CONFIG)

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(password: str, hashed: str) -> bool:
    return pwd_context.verify(password, hashed)

def create_token(user_id: int, is_admin: bool = False) -> str:
    payload = {
        "sub": str(user_id),
        "adm": is_admin,
        "exp": datetime.utcnow() + timedelta(days=TOKEN_EXPIRE_DAYS)
    }
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(bearer)) -> int:
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        return int(payload["sub"])
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

def get_admin_user(credentials: HTTPAuthorizationCredentials = Depends(bearer)) -> int:
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        if not payload.get("adm"):
            raise HTTPException(status_code=403, detail="Admin access required")
        return int(payload["sub"])
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

# ── Auth endpoints ──

@app.post("/auth/register")
@limiter.limit("5/hour")
def register(request: Request, req: AuthRequest):
    if not req.invite_code:
        raise HTTPException(status_code=400, detail="邀请码不能为空")
    conn = get_db()
    cur = conn.cursor()
    try:
        # 验证邀请码
        cur.execute(
            "SELECT id FROM invite_codes WHERE code=%s AND is_active=TRUE AND used_by IS NULL",
            (req.invite_code,)
        )
        code_row = cur.fetchone()
        if not code_row:
            raise HTTPException(status_code=400, detail="邀请码无效或已被使用")

        cur.execute("SELECT id FROM users WHERE email=%s", (req.email,))
        if cur.fetchone():
            raise HTTPException(status_code=400, detail="Email already registered")

        cur.execute(
            "INSERT INTO users (email, password_hash) VALUES (%s, %s) RETURNING id, is_admin",
            (req.email, hash_password(req.password))
        )
        row = cur.fetchone()
        user_id = row[0]

        # 标记邀请码已使用
        cur.execute(
            "UPDATE invite_codes SET used_by=%s, used_at=NOW() WHERE id=%s",
            (user_id, code_row[0])
        )
        conn.commit()
        return {"token": create_token(user_id, row[1]), "email": req.email, "is_admin": row[1]}
    finally:
        cur.close(); conn.close()

@app.post("/auth/login")
@limiter.limit("10/minute")
def login(request: Request, req: AuthRequest):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id, password_hash, is_admin FROM users WHERE email=%s", (req.email,))
        row = cur.fetchone()
        if not row or not verify_password(req.password, row[1]):
            raise HTTPException(status_code=401, detail="Invalid email or password")
        return {"token": create_token(row[0], row[2]), "email": req.email, "is_admin": row[2]}
    finally:
        cur.close(); conn.close()

@app.patch("/auth/change-password")
def change_password(req: ChangePasswordRequest, user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT password_hash FROM users WHERE id=%s", (user_id,))
        row = cur.fetchone()
        if not row or not verify_password(req.current_password, row[0]):
            raise HTTPException(status_code=400, detail="Current password is incorrect")
        cur.execute("UPDATE users SET password_hash=%s WHERE id=%s", (hash_password(req.new_password), user_id))
        conn.commit()
        return {"success": True}
    finally:
        cur.close(); conn.close()

# ── Application endpoints (auth required) ──

@app.get("/applications")
def get_applications(user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""
        SELECT id, company, position, applied_date, location, link, feedback, work_type, notes
        FROM job_applications
        WHERE user_id=%s
        ORDER BY applied_date DESC NULLS LAST, id DESC
    """, (user_id,))
    rows = [dict(r) for r in cur.fetchall()]
    cur.close(); conn.close()
    for r in rows:
        if r['applied_date']:
            r['applied_date'] = r['applied_date'].isoformat()
    return rows

@app.post("/applications")
def add_application(req: ApplicationRequest, user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("""
            INSERT INTO job_applications (company, position, applied_date, location, link, feedback, work_type, notes, user_id)
            VALUES (%s, %s, %s::date, %s, %s, %s, %s, %s, %s)
        """, (req.company, req.position, req.applied_date or None,
              req.location, req.link, req.feedback, req.work_type, req.notes, user_id))
        conn.commit()
        return {"success": True}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cur.close(); conn.close()

@app.put("/applications/{app_id}")
def update_application(app_id: int, req: ApplicationRequest, user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("""
            UPDATE job_applications
            SET company=%s, position=%s, applied_date=%s::date,
                location=%s, link=%s, feedback=%s, work_type=%s, notes=%s
            WHERE id=%s AND user_id=%s
        """, (req.company, req.position, req.applied_date or None,
              req.location, req.link, req.feedback, req.work_type, req.notes, app_id, user_id))
        conn.commit()
        return {"success": True}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cur.close(); conn.close()

@app.delete("/applications/{app_id}")
def delete_application(app_id: int, user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM job_applications WHERE id=%s AND user_id=%s", (app_id, user_id))
        conn.commit()
        return {"success": True}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cur.close(); conn.close()

# ── Stats endpoints (auth required) ──

@app.get("/stats/summary")
def stats_summary(user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        cur.execute("""
            SELECT
                COUNT(*) as total,
                COUNT(*) FILTER (WHERE feedback IS NULL) as pending,
                COUNT(DISTINCT location) as countries
            FROM job_applications WHERE user_id=%s
        """, (user_id,))
        return dict(cur.fetchone())
    finally:
        cur.close(); conn.close()

@app.get("/stats/countries")
def stats_countries(user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        cur.execute("""
            SELECT location, COUNT(*) as count
            FROM job_applications
            WHERE user_id=%s AND location IS NOT NULL AND location != '' AND location != 'NaN'
            GROUP BY location ORDER BY count DESC LIMIT 5
        """, (user_id,))
        return [dict(r) for r in cur.fetchall()]
    finally:
        cur.close(); conn.close()

@app.get("/stats/worktype")
def stats_worktype(user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        cur.execute("""
            SELECT
                COUNT(*) FILTER (WHERE work_type = 'Remote') as remote,
                COUNT(*) FILTER (WHERE work_type = 'Onsite') as onsite,
                COUNT(*) FILTER (WHERE work_type = 'Hybrid') as hybrid
            FROM job_applications WHERE user_id=%s
        """, (user_id,))
        return dict(cur.fetchone())
    finally:
        cur.close(); conn.close()

# ── Admin endpoints ──

@app.get("/admin/stats")
def admin_stats(admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT COUNT(*) FROM users")
        total_users = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM users WHERE created_at::date = CURRENT_DATE")
        new_today = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM user_feedback")
        total_feedback = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM invite_codes WHERE is_active=TRUE AND used_by IS NULL")
        available_invites = cur.fetchone()[0]
        return {
            "total_users": total_users,
            "new_today": new_today,
            "total_feedback": total_feedback,
            "available_invites": available_invites
        }
    finally:
        cur.close(); conn.close()



@app.get("/admin/users")
def admin_list_users(admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        cur.execute("""
            SELECT id, email, is_admin, created_at
            FROM users ORDER BY id
        """)
        rows = [dict(r) for r in cur.fetchall()]
        for r in rows:
            if r['created_at']:
                r['created_at'] = r['created_at'].isoformat()
        return rows
    finally:
        cur.close(); conn.close()

@app.delete("/admin/users/{uid}")
def admin_delete_user(uid: int, admin_id: int = Depends(get_admin_user)):
    if uid == admin_id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("DELETE FROM job_applications WHERE user_id=%s", (uid,))
        cur.execute("DELETE FROM users WHERE id=%s", (uid,))
        conn.commit()
        return {"success": True}
    finally:
        cur.close(); conn.close()

@app.patch("/admin/users/{uid}/toggle-admin")
def admin_toggle_admin(uid: int, admin_id: int = Depends(get_admin_user)):
    if uid == admin_id:
        raise HTTPException(status_code=400, detail="Cannot change your own admin status")
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("UPDATE users SET is_admin = NOT is_admin WHERE id=%s RETURNING is_admin", (uid,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        new_status = row[0]
        conn.commit()
        return {"is_admin": new_status}
    finally:
        cur.close(); conn.close()

@app.patch("/admin/users/{uid}/reset-password")
def admin_reset_password(uid: int, req: ResetPasswordRequest, admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("UPDATE users SET password_hash=%s WHERE id=%s", (hash_password(req.new_password), uid))
        conn.commit()
        return {"success": True}
    finally:
        cur.close(); conn.close()

# ── Admin: invite codes ──

@app.post("/admin/invite-codes")
def admin_create_invite(admin_id: int = Depends(get_admin_user)):
    import secrets as _secrets
    code = _secrets.token_urlsafe(8)
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO invite_codes (code, created_by) VALUES (%s, %s) RETURNING id, code, created_at",
            (code, admin_id)
        )
        row = cur.fetchone()
        conn.commit()
        return {"id": row[0], "code": row[1], "created_at": row[2].isoformat()}
    finally:
        cur.close(); conn.close()

@app.get("/admin/invite-codes")
def admin_list_invites(admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("""
        SELECT ic.id, ic.code, ic.is_active, ic.created_at,
               ic.used_at, u.email as used_by_email
        FROM invite_codes ic
        LEFT JOIN users u ON u.id = ic.used_by
        ORDER BY ic.created_at DESC
    """)
    rows = [dict(r) for r in cur.fetchall()]
    cur.close(); conn.close()
    for r in rows:
        if r['created_at']: r['created_at'] = r['created_at'].isoformat()
        if r['used_at']:    r['used_at']    = r['used_at'].isoformat()
    return rows

@app.delete("/admin/invite-codes/{code_id}")
def admin_revoke_invite(code_id: int, admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("UPDATE invite_codes SET is_active=FALSE WHERE id=%s AND used_by IS NULL", (code_id,))
        conn.commit()
        return {"success": True}
    finally:
        cur.close(); conn.close()

# ── Feedback endpoints ──

@app.post("/feedback")
def submit_feedback(req: FeedbackRequest, user_id: int = Depends(get_current_user)):
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO user_feedback (user_id, category, content) VALUES (%s, %s, %s)",
            (user_id, req.category, req.content)
        )
        conn.commit()
        return {"ok": True}
    finally:
        cur.close(); conn.close()

@app.get("/admin/feedback")
def get_all_feedback(admin_id: int = Depends(get_admin_user)):
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        cur.execute("""
            SELECT uf.id, u.email, uf.category, uf.content, uf.created_at
            FROM user_feedback uf
            JOIN users u ON u.id = uf.user_id
            ORDER BY uf.created_at DESC
        """)
        rows = []
        for r in cur.fetchall():
            row = dict(r)
            if row.get("created_at"):
                row["created_at"] = row["created_at"].isoformat()
            rows.append(row)
        return rows
    finally:
        cur.close(); conn.close()

# ── Image parse endpoint ──

PARSE_IMAGE_PROMPT = """Extract job application information from this image.
Return ONLY a JSON object with these exact keys (use null for any missing or unclear field):
{
  "company": "company or organization name",
  "position": "job title or role",
  "applied_date": "date in YYYY-MM-DD format or null",
  "location": "country or region only (e.g. Singapore, Netherlands, United States)",
  "link": "full URL including https:// if any link is visible, otherwise null",
  "work_type": "Remote, Onsite, or Hybrid or null",
  "feedback": "application status if shown (e.g. Interview, Offer, Fail, Online Assessment) or null"
}
Return only the JSON object, no markdown, no explanation."""

@app.post("/applications/parse-image")
async def parse_image(file: UploadFile = File(...), user_id: int = Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=501, detail="ANTHROPIC_API_KEY not configured")
    image_data = await file.read()
    if len(image_data) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Image too large (max 10MB)")
    b64 = base64.standard_b64encode(image_data).decode("utf-8")
    media_type = file.content_type or "image/jpeg"
    if media_type not in ("image/jpeg", "image/png", "image/gif", "image/webp"):
        media_type = "image/jpeg"
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=512,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                    {"type": "text", "text": PARSE_IMAGE_PROMPT}
                ]
            }]
        )
        text = response.content[0].text.strip()
        if text.startswith("```"):
            text = re.sub(r"^```[a-z]*\n?", "", text).rstrip("`").strip()
        return json.loads(text)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="AI returned unparseable response")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── Chat endpoint ──

@app.post("/chat")
@limiter.limit("30/minute")
def chat(request: Request, req: ChatRequest, user_id: int = Depends(get_current_user)):
    # 每日调用限制
    today = datetime.utcnow().date()
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO chat_usage (user_id, date, count) VALUES (%s, %s, 1) "
            "ON CONFLICT (user_id, date) DO UPDATE SET count = chat_usage.count + 1 "
            "RETURNING count",
            (user_id, today)
        )
        usage = cur.fetchone()[0]
        conn.commit()
    finally:
        cur.close(); conn.close()
    if usage > CHAT_DAILY_LIMIT:
        eng = _is_english(req.message)
        detail = (f"Daily limit reached ({CHAT_DAILY_LIMIT} queries). Come back tomorrow!"
                  if eng else f"今日提问已达上限（{CHAT_DAILY_LIMIT} 次），明天再来吧")
        raise HTTPException(status_code=429, detail=detail)

    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

    # Step 1: NL → SQL
    system_prompt = SQL_SYSTEM_PROMPT.replace("{user_id}", str(user_id))
    history = [{"role": m.role, "content": m.content} for m in (req.history or [])]
    messages = [
        {"role": "system", "content": system_prompt},
        *history[-6:],
        {"role": "user", "content": req.message}
    ]
    sql_resp = client.chat.completions.create(
        model="deepseek-v4-flash", messages=messages, temperature=0
    )
    raw = sql_resp.choices[0].message.content.strip()

    # 拒绝语直接返回
    if not raw.upper().lstrip().startswith("SELECT"):
        return {"answer": raw, "sql": None}

    # 提取第一条 SELECT 语句：去掉末尾分号及之后的多余内容
    sql_or_reject = raw.split(";")[0].strip()

    if not sql_or_reject.upper().startswith("SELECT"):
        return {"answer": sql_or_reject, "sql": None}

    # Step 2: 执行 SQL（安全检查）
    err = validate_chat_sql(sql_or_reject, user_id)
    if err:
        eng = _is_english(req.message)
        msg = f"Generated query was blocked: {err}" if eng else f"生成的查询已被拦截：{err}"
        return {"answer": msg, "sql": None}
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SET LOCAL statement_timeout = '5000'")  # 5s 超时
        cur.execute(sql_or_reject)
        rows = [dict(r) for r in cur.fetchall()]
        cur.close(); conn.close()
        for r in rows:
            for k, v in r.items():
                if hasattr(v, 'isoformat'):
                    r[k] = v.isoformat()
    except Exception as e:
        eng = _is_english(req.message)
        msg = f"Query error: {str(e)}" if eng else f"查询出错：{str(e)}"
        return {"answer": msg, "sql": sql_or_reject}

    # Step 3: 结果 → 自然语言
    eng = _is_english(req.message)
    explain_system = EXPLAIN_SYSTEM_PROMPT_EN if eng else EXPLAIN_SYSTEM_PROMPT_ZH
    if eng:
        user_content = (f"Answer in English only.\n\n"
                        f"Question: {req.message}\n\nQuery result: {rows}")
    else:
        user_content = f"问题：{req.message}\n\n查询结果：{rows}"
    explain_messages = [
        {"role": "system", "content": explain_system},
        {"role": "user", "content": user_content}
    ]
    explain_resp = client.chat.completions.create(
        model="deepseek-v4-flash", messages=explain_messages, temperature=0
    )
    answer = explain_resp.choices[0].message.content.strip()
    return {"answer": answer, "sql": sql_or_reject}

# ── Analyze endpoint (Claude Sonnet 4.6) ──

ANALYZE_DAILY_LIMIT = 100

@app.post("/analyze")
@limiter.limit("10/minute")
def analyze(request: Request, req: AnalyzeRequest, user_id: int = Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=501, detail="ANTHROPIC_API_KEY not configured")

    today = datetime.utcnow().date()
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO chat_usage (user_id, date, count) VALUES (%s, %s, 1) "
            "ON CONFLICT (user_id, date) DO UPDATE SET count = chat_usage.count + 1 "
            "RETURNING count",
            (user_id, today)
        )
        usage = cur.fetchone()[0]
        conn.commit()
    finally:
        cur.close(); conn.close()

    if usage > ANALYZE_DAILY_LIMIT:
        raise HTTPException(status_code=429, detail="今日分析次数已达上限，明天再来吧")

    if req.type == "fate":
        ds_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
        resp = ds_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": req.message}],
            temperature=0.9,
            max_tokens=512,
        )
        return {"answer": resp.choices[0].message.content.strip()}

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": req.message}]
    )
    return {"answer": response.content[0].text.strip()}

# ── Visual resume optimization endpoint ──

class VisualResumeRequest(BaseModel):
    images: List[str]  # base64 JPEG, max 3 pages
    jd: str = ""

@app.post("/optimize-resume-visual")
@limiter.limit("5/minute")
def optimize_resume_visual(request: Request, req: VisualResumeRequest, user_id: int = Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=501, detail="ANTHROPIC_API_KEY not configured")

    today = datetime.utcnow().date()
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO chat_usage (user_id, date, count) VALUES (%s, %s, 1) "
            "ON CONFLICT (user_id, date) DO UPDATE SET count = chat_usage.count + 1 "
            "RETURNING count",
            (user_id, today)
        )
        usage = cur.fetchone()[0]
        conn.commit()
    finally:
        cur.close(); conn.close()

    if usage > ANALYZE_DAILY_LIMIT:
        raise HTTPException(status_code=429, detail="今日分析次数已达上限，明天再来吧")

    images = req.images[:3]
    content = []
    for i, img_b64 in enumerate(images):
        content.append({"type": "text", "text": f"Resume page {i + 1}:"})
        content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}
        })

    jd_section = f"\n\n[Job Description]\n{req.jd}" if req.jd.strip() else ""
    content.append({"type": "text", "text": f"""Analyze the visual design of this resume and generate an optimized version as a complete, self-contained HTML document.

Requirements:
- Replicate the exact visual design: layout, color scheme, typography, spacing, section structure
- Preserve all factual content: company names, job titles, dates, schools, degrees — never fabricate
- Optimize wording with strong action verbs and ATS keywords woven in naturally — no keyword stuffing
- Keep the same section order and number of bullet points per role as the original
- Do not invent numbers or percentages not present in the original
- Match the language of the original resume exactly
- Keep total content length similar to the original — do not significantly expand
- HTML must be completely self-contained: all CSS in a <style> tag, no external dependencies
- Must be print-ready: A4 page size, proper margins for clean PDF export via browser print
- Highlight every word or phrase you changed or added by wrapping it with <span class="opt-highlight">...</span>. Add this CSS: .opt-highlight {{ background: #fef08a; border-radius: 2px; }}{jd_section}

Output ONLY the complete HTML document starting with <!DOCTYPE html>. No explanations, no markdown code blocks."""})

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        messages=[{"role": "user", "content": content}]
    )

    html = response.content[0].text.strip()
    html = re.sub(r'^```[a-z]*\n?', '', html)
    html = re.sub(r'\n?```$', '', html)

    return {"html": html}

# ── Word resume optimization endpoint ──

class WordResumeRequest(BaseModel):
    word_b64: str
    jd: str = ""

@app.post("/optimize-word-resume")
@limiter.limit("5/minute")
def optimize_word_resume(request: Request, req: WordResumeRequest, user_id: int = Depends(get_current_user)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=501, detail="ANTHROPIC_API_KEY not configured")

    today = datetime.utcnow().date()
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO chat_usage (user_id, date, count) VALUES (%s, %s, 1) "
            "ON CONFLICT (user_id, date) DO UPDATE SET count = chat_usage.count + 1 "
            "RETURNING count",
            (user_id, today)
        )
        usage = cur.fetchone()[0]
        conn.commit()
    finally:
        cur.close(); conn.close()

    if usage > ANALYZE_DAILY_LIMIT:
        raise HTTPException(status_code=429, detail="今日分析次数已达上限，明天再来吧")

    try:
        docx_bytes = base64.b64decode(req.word_b64)
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        word_html = result.value
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Word 解析失败: {str(e)}")

    jd_section = f"\n\n[Job Description]\n{req.jd}" if req.jd.strip() else ""
    prompt = f"""You are a professional resume designer and career consultant. Below is the HTML source of my resume (converted from Word), preserving the original formatting, colors, fonts, and structure.

Generate a complete, self-contained HTML document that:
- Replicates the visual design from the source HTML (colors, fonts, layout, spacing, section structure)
- Preserves all factual content: company names, job titles, dates, schools, degrees — never fabricate or alter them
- Optimizes wording with strong action verbs and ATS keywords woven in naturally — no keyword stuffing
- Keeps the same section order and number of bullet points per role as the original
- Does not invent numbers or percentages not present in the original
- Matches the language of the original resume exactly
- Keeps total content length similar to the original — do not significantly expand
- HTML must be completely self-contained: all CSS in a <style> tag, no external dependencies
- Must be print-ready: A4 page size, proper margins for PDF export via browser print
- Highlight every word or phrase you changed or added by wrapping it with <span class="opt-highlight">...</span>. Add this CSS: .opt-highlight {{ background: #fef08a; border-radius: 2px; }}

[Original Resume HTML]
{word_html}{jd_section}

Output ONLY the complete HTML document starting with <!DOCTYPE html>. No explanations, no markdown code blocks."""

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}]
    )

    html = response.content[0].text.strip()
    html = re.sub(r'^```[a-z]*\n?', '', html)
    html = re.sub(r'\n?```$', '', html)
    return {"html": html}

# ── Resume export helpers ──

def _find_cjk_font():
    """Find an available CJK font on the system."""
    candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None

def _parse_markdown_lines(content: str):
    """Parse markdown content into structured lines for document generation."""
    lines = content.strip().split('\n')
    parsed = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            parsed.append(('hr', None))
            i += 1
            continue

        # Heading
        h_match = re.match(r'^(#{1,3})\s+(.+)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            parsed.append(('heading', (level, text)))
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^(\s*)[-*•]\s+(.+)', line)
        if bullet_match:
            items = []
            while i < len(lines):
                bm = re.match(r'^(\s*)[-*•]\s+(.+)', lines[i])
                if bm:
                    items.append(bm.group(2))
                    i += 1
                else:
                    break
            parsed.append(('bullet_list', items))
            continue

        # Numbered list
        num_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if num_match:
            items = []
            while i < len(lines):
                nm = re.match(r'^(\s*)\d+\.\s+(.+)', lines[i])
                if nm:
                    items.append(nm.group(2))
                    i += 1
                else:
                    break
            parsed.append(('numbered_list', items))
            continue

        # Paragraph (skip empty lines, collect consecutive non-markup lines)
        if line.strip():
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not re.match(
                r'^(#{1,3}\s|[-*•]\s|\d+\.\s|---|\*\*\*|___)', lines[i]
            ):
                para_lines.append(lines[i])
                i += 1
            parsed.append(('paragraph', ' '.join(para_lines)))
        else:
            i += 1

    return parsed

def _split_inline_bold(text: str):
    """Split text into (type, content) segments, tagging bold portions."""
    pattern = re.compile(r'\*\*(.+?)\*\*')
    result = []
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            result.append(('text', text[last_end:m.start()]))
        result.append(('bold', m.group(1)))
        last_end = m.end()
    if last_end < len(text):
        result.append(('text', text[last_end:]))
    return result if result else [('text', text)]


# ── Export endpoint ──

@app.post("/export-resume")
def export_resume(req: ExportRequest, user_id: int = Depends(get_current_user)):
    content = req.content
    fmt = req.format.lower()

    if not content or not content.strip():
        raise HTTPException(status_code=400, detail="Content is empty")
    if fmt not in ('docx', 'pdf'):
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")

    parsed = _parse_markdown_lines(content)

    if fmt == 'docx':
        return _generate_docx(parsed)
    else:
        return _generate_pdf(parsed)


def _generate_docx(parsed):
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for kind, data in parsed:
        if kind == 'heading':
            level, text = data
            h = doc.add_heading(text, level=min(level, 3))
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

        elif kind == 'hr':
            doc.add_paragraph('─' * 60)

        elif kind == 'paragraph':
            _add_docx_para(doc, data)

        elif kind == 'bullet_list':
            for item in data:
                para = doc.add_paragraph()
                para.style = doc.styles['List Bullet'] if 'List Bullet' in [s.name for s in doc.styles] else doc.styles['Normal']
                _add_formatted_runs(para, item)

        elif kind == 'numbered_list':
            for idx, item in enumerate(data, 1):
                para = doc.add_paragraph()
                _add_formatted_runs(para, item, prefix=f'{idx}. ')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename="optimized-resume.docx"'}
    )


def _add_docx_para(doc, text: str):
    para = doc.add_paragraph()
    _add_formatted_runs(para, text)


def _add_formatted_runs(para, text: str, prefix: str = ''):
    if prefix:
        para.add_run(prefix)
    for seg_type, seg_text in _split_inline_bold(text):
        run = para.add_run(seg_text)
        if seg_type == 'bold':
            run.bold = True


def _generate_pdf(parsed):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    cjk_font_path = _find_cjk_font()
    if cjk_font_path:
        pdf.add_font('CJK', '', cjk_font_path, uni=True)
        pdf.add_font('CJK', 'B', cjk_font_path, uni=True)
        font_name = 'CJK'
        has_bold = True
    else:
        font_name = 'Helvetica'
        has_bold = False

    for kind, data in parsed:
        if kind == 'heading':
            level, text = data
            sizes = {1: 16, 2: 13, 3: 11.5}
            pdf.set_font(font_name, 'B' if has_bold else '', sizes.get(level, 13))
            pdf.ln(3)
            pdf.multi_cell(0, 7, text)
            pdf.ln(2)

        elif kind == 'hr':
            pdf.ln(3)
            y = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(10, y, 200, y)
            pdf.ln(4)

        elif kind == 'paragraph':
            pdf.set_font(font_name, '', 10.5)
            pdf.ln(1)
            _pdf_write_formatted(pdf, data, font_name, has_bold)
            pdf.ln(2)

        elif kind == 'bullet_list':
            pdf.set_font(font_name, '', 10.5)
            for item in data:
                pdf.cell(6, 6, '•')
                _pdf_write_formatted(pdf, item, font_name, has_bold)
                pdf.ln(2)
            pdf.ln(1)

        elif kind == 'numbered_list':
            pdf.set_font(font_name, '', 10.5)
            for idx, item in enumerate(data, 1):
                pdf.cell(8, 6, f'{idx}.')
                _pdf_write_formatted(pdf, item, font_name, has_bold)
                pdf.ln(2)
            pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type='application/pdf',
        headers={'Content-Disposition': 'attachment; filename="optimized-resume.pdf"'}
    )


def _pdf_write_formatted(pdf, text: str, font_name: str, has_bold: bool):
    """Write text with inline bold to PDF using write() for proper inline flow."""
    segments = _split_inline_bold(text)
    for seg_type, seg_text in segments:
        if seg_type == 'bold' and has_bold:
            pdf.set_font(font_name, 'B', 10.5)
        else:
            pdf.set_font(font_name, '', 10.5)
        pdf.write(6, seg_text)
    pdf.ln()

# ── Public endpoints ──

@app.get("/rss-proxy")
def rss_proxy(url: str, user_id: int = Depends(get_current_user)):
    from urllib.parse import urlparse
    import urllib.request
    ALLOWED_HOSTS = {
        'www.anthropic.com',
        'www.theverge.com',
        'hnrss.org',
        'www.latent.space',
        'www.technologyreview.com',
    }
    parsed = urlparse(url)
    if parsed.hostname not in ALLOWED_HOSTS or parsed.scheme not in ('http', 'https'):
        raise HTTPException(status_code=400, detail="URL not allowed")
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (compatible; JobTrackBot/1.0)'})
        with urllib.request.urlopen(req, timeout=12) as resp:
            content = resp.read()
        return Response(content=content, media_type="application/xml; charset=utf-8")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch feed: {e}")

@app.get("/health")
def health():
    return {"status": "ok"}
